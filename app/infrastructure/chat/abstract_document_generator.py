from abc import ABC, abstractmethod
from io import BytesIO

from docx import Document

from app.domain.models.chat_response import ChatResponse


class AbstractDocumentGenerator(ABC):
    @abstractmethod
    def generate(self, chat_response: ChatResponse, file_stream: BytesIO):
        pass


class DocxGenerator(AbstractDocumentGenerator):

    def generate(self, chat_response: ChatResponse, file_stream: BytesIO):
        document = Document()

        # Add probability and description
        document.add_heading("Candidate Profile", level=1)
        document.add_paragraph(f"Probability: {round(chat_response.probability*100)}%")
        document.add_paragraph(chat_response.probability_description)

        # Add experiences
        document.add_heading("Experiences", level=1)
        for exp in chat_response.experiences:
            document.add_heading(exp.company_name, level=2)
            document.add_paragraph(f"{exp.title}")
            document.add_heading("Highlights:", level=3)
            for highlight in exp.highlights:
                document.add_paragraph(f"- {highlight}")

        # Add cover letter text
        document.add_heading("Cover Letter", level=1)
        document.add_paragraph(chat_response.cover_letter_text)

        # Save the document
        document.save(file_stream)
